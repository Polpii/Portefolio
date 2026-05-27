"""Add horizontal underline (pict) to RESEARCH EXPERIENCE, PUBLICATIONS, ACTIVITIES headers
by cloning the existing pict from the SKILLS header."""
from docx import Document
from lxml import etree
from copy import deepcopy
import subprocess

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document('public/resume_2026.docx')

# Locate source pict (inside SKILLS paragraph)
src_para = None
for p in doc.paragraphs:
    if p.text.strip() == 'SKILLS':
        src_para = p
        break

# Find the run containing the pict
src_pict_run = None
for r in src_para._p.findall(f'{W}r'):
    if r.find(f'.//{W}pict') is not None:
        src_pict_run = r
        break

assert src_pict_run is not None, "no pict run found in SKILLS"
print(f"Source run XML length: {len(etree.tostring(src_pict_run))}")

targets = {'RESEARCH EXPERIENCE', 'PUBLICATIONS', 'ACTIVITIES'}
for p in doc.paragraphs:
    t = p.text.strip()
    if t in targets:
        # Skip if already has pict
        if p._p.findall(f'.//{W}pict'):
            print(f"  {t}: already has pict, skipping")
            continue
        # Clone source run and insert as first child of paragraph (before existing runs)
        new_run = deepcopy(src_pict_run)
        # We need to make sure each pict has unique IDs to avoid Word complaints
        # But docx2pdf/Word usually tolerates duplicates; we'll rename anchorId & ids
        import re, uuid
        xml_str = etree.tostring(new_run).decode()
        # Replace anchorId values, o:spid, and group id with unique values
        new_anchor = uuid.uuid4().hex[:8].upper()
        xml_str = re.sub(r'w14:anchorId="[^"]*"', f'w14:anchorId="{new_anchor}"', xml_str)
        new_run = etree.fromstring(xml_str)
        # Insert before existing runs (after pPr)
        pPr = p._p.find(f'{W}pPr')
        if pPr is not None:
            pPr.addnext(new_run)
        else:
            p._p.insert(0, new_run)
        print(f"  {t}: pict inserted")

doc.save('public/resume_2026.docx')
print("Saved docx")

# Regenerate PDF
from docx2pdf import convert
convert('public/resume_2026.docx', 'private-docs/resume_2026.pdf')
print("Saved PDF")
