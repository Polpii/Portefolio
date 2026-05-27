"""Clean approach: remove all VML pict horizontal lines and add a uniform
paragraph bottom border (pBdr) to all 6 section headers."""
from docx import Document
from lxml import etree
from copy import deepcopy
import re

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
doc = Document('public/resume_2026.docx')

# 1) Remove all pict-bearing runs across the whole document
removed = 0
for p in doc.paragraphs:
    for r in list(p._p.findall(f'{W}r')):
        if r.find(f'.//{W}pict') is not None:
            r.getparent().remove(r)
            removed += 1
print(f"Removed {removed} pict runs")

# 2) Add a bottom border to each section header paragraph
HEADERS = {'RESEARCH EXPERIENCE','PUBLICATIONS','EDUCATION','SKILLS','WORK EXPERIENCE','ACTIVITIES'}

def make_pBdr():
    pBdr = etree.SubElement(etree.Element(f'{W}tmp'), f'{W}pBdr')
    bottom = etree.SubElement(pBdr, f'{W}bottom')
    bottom.set(f'{W}val', 'single')
    bottom.set(f'{W}sz', '6')        # 0.75pt
    bottom.set(f'{W}space', '1')
    bottom.set(f'{W}color', 'auto')
    return pBdr

for p in doc.paragraphs:
    if p.text.strip() in HEADERS:
        pPr = p._p.find(f'{W}pPr')
        if pPr is None:
            pPr = etree.SubElement(p._p, f'{W}pPr')
            p._p.insert(0, pPr)
        # Remove any existing pBdr
        existing = pPr.find(f'{W}pBdr')
        if existing is not None:
            pPr.remove(existing)
        # Insert new pBdr — must be in correct schema order; place after pStyle
        pBdr = make_pBdr()
        # Find pStyle to insert after it
        pStyle = pPr.find(f'{W}pStyle')
        if pStyle is not None:
            pStyle.addnext(pBdr)
        else:
            pPr.insert(0, pBdr)
        print(f"Added bottom border: {p.text.strip()}")

doc.save('public/resume_2026.docx')
print("Saved docx")

from docx2pdf import convert
convert('public/resume_2026.docx', 'private-docs/resume_2026.pdf')
print("Saved PDF")
