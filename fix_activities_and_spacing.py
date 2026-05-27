"""Fix resume:
1. Remove (GPT-4 class) from skills line.
2. Ensure exactly one blank paragraph before each section header.
3. Replace ACTIVITIES content with structured entries matching WORK EXPERIENCE style.
"""
from docx import Document
from lxml import etree
from copy import deepcopy

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
Wq = '{' + W + '}'
nsmap = {'w': W}

doc = Document('public/resume_2026.docx')

# ---------- 1) Remove "(GPT-4 class)" ----------
for p in doc.paragraphs:
    for r in p._p.findall(f'{Wq}r'):
        for t in r.findall(f'{Wq}t'):
            if t.text and 'GPT-4 class' in t.text:
                t.text = t.text.replace(' (GPT-4 class)', '').replace('(GPT-4 class)', '')

# ---------- 2) ACTIVITIES section rebuild ----------
HEADERS = {'RESEARCH EXPERIENCE','PUBLICATIONS','EDUCATION','SKILLS','WORK EXPERIENCE','ACTIVITIES'}

def w(tag):
    return etree.SubElement.__self_class__  # unused

def make_company_line(company, location):
    p = etree.fromstring(f'''<w:p xmlns:w="{W}">
      <w:pPr>
        <w:tabs><w:tab w:val="right" w:pos="9826"/></w:tabs>
        <w:spacing w:before="160"/>
        <w:ind w:left="14"/>
        <w:rPr><w:sz w:val="18"/></w:rPr>
      </w:pPr>
      <w:r><w:rPr><w:b/><w:w w:val="120"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{company}</w:t></w:r>
      <w:r><w:rPr><w:b/><w:sz w:val="18"/></w:rPr><w:tab/></w:r>
      <w:r><w:rPr><w:w w:val="110"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{location}</w:t></w:r>
    </w:p>''')
    return p

def make_role_line(role, dates):
    p = etree.fromstring(f'''<w:p xmlns:w="{W}">
      <w:pPr>
        <w:tabs><w:tab w:val="right" w:pos="9826"/></w:tabs>
        <w:spacing w:before="1"/>
        <w:ind w:left="14"/>
        <w:rPr><w:i/><w:sz w:val="18"/></w:rPr>
      </w:pPr>
      <w:r><w:rPr><w:b/><w:i/><w:w w:val="110"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{role}</w:t></w:r>
      <w:r><w:rPr><w:b/><w:i/><w:sz w:val="18"/></w:rPr><w:tab/></w:r>
      <w:r><w:rPr><w:i/><w:w w:val="110"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{dates}</w:t></w:r>
    </w:p>''')
    return p

def make_bullet(text):
    # Escape XML-special chars in text
    text = (text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'))
    p = etree.fromstring(f'''<w:p xmlns:w="{W}">
      <w:pPr>
        <w:pStyle w:val="Paragraphedeliste"/>
        <w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>
        <w:tabs><w:tab w:val="left" w:pos="733"/></w:tabs>
        <w:ind w:left="733" w:hanging="359"/>
        <w:rPr><w:sz w:val="18"/></w:rPr>
      </w:pPr>
      <w:r><w:rPr><w:w w:val="110"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>
    </w:p>''')
    return p

def make_plain_line(text):
    text = (text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'))
    p = etree.fromstring(f'''<w:p xmlns:w="{W}">
      <w:pPr>
        <w:spacing w:before="80"/>
        <w:ind w:left="14"/>
        <w:rPr><w:sz w:val="18"/></w:rPr>
      </w:pPr>
      <w:r><w:rPr><w:b/><w:w w:val="110"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{text.split(":",1)[0]}:</w:t></w:r>
      <w:r><w:rPr><w:w w:val="110"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{text.split(":",1)[1] if ":" in text else ""}</w:t></w:r>
    </w:p>''')
    return p

# Find ACTIVITIES header and locate its content paragraphs (everything after until end / next header)
body = doc.element.body
paragraphs = doc.paragraphs

# Find ACTIVITIES index
activities_idx = None
for i, p in enumerate(paragraphs):
    if p.text.strip() == 'ACTIVITIES':
        activities_idx = i
        break

assert activities_idx is not None
activities_p = paragraphs[activities_idx]._p

# Remove all paragraphs after ACTIVITIES (they are the old chess one-liner + any trailing)
# Collect siblings after ACTIVITIES (within sectPr boundaries)
to_remove = []
sib = activities_p.getnext()
while sib is not None:
    if sib.tag == f'{Wq}p':
        to_remove.append(sib)
        sib = sib.getnext()
    else:
        # sectPr or other — stop
        break

# But preserve final sectPr-bearing paragraph if any
# Inspect last paragraph in to_remove: if it contains sectPr inside pPr, keep it
# (Word stores sectPr in the last paragraph's pPr)
keep_last = None
if to_remove:
    last = to_remove[-1]
    pPr = last.find(f'{Wq}pPr')
    if pPr is not None and pPr.find(f'{Wq}sectPr') is not None:
        keep_last = to_remove.pop()  # don't remove it; we'll clear its runs but keep sectPr

for p in to_remove:
    p.getparent().remove(p)
print(f"Removed {len(to_remove)} old ACTIVITIES paragraph(s)")

# Build new content
new_paras = [
    make_company_line('LCE Chess club', 'La Ciotat, FR'),
    make_role_line('Chess instructor', '2007 – 2018'),
    make_bullet('DIFF certificate (chess instructor) — French Chess Federation'),
    make_bullet('FIDE Elo rating: 1880 (World Chess Federation)'),
    make_bullet('Instructed a senior class of 10 students'),
    make_bullet('2nd place — PACA Open; 1st place — National IV Team Competition'),
    make_company_line('Eagles — ESILV US Football Team', 'Paris, FR'),
    make_role_line('Wide Receiver', '2017 – 2018'),
    make_bullet('University French Champion (2017) — US Football Association'),
    make_plain_line('Sports: Rugby, Skiing, Swimming, Krav Maga'),
    make_plain_line('Music: composition with FL Studio (Fruity Loops)'),
]

# Insert after ACTIVITIES header (before keep_last if exists, else append)
anchor = activities_p
for np in new_paras:
    anchor.addnext(np)
    anchor = np

# If keep_last exists, ensure it stays at the end (it was already preserved)

# ---------- 3) Ensure exactly one blank paragraph before each section header ----------
def is_blank_para(p):
    # No text
    for t in p.findall(f'.//{Wq}t'):
        if t.text and t.text.strip():
            return False
    return True

def make_blank_para():
    return etree.fromstring(f'<w:p xmlns:w="{W}"><w:pPr><w:pStyle w:val="Corpsdetexte"/><w:rPr><w:sz w:val="18"/></w:rPr></w:pPr></w:p>')

# Re-fetch paragraphs after modification
for p_obj in list(doc.paragraphs):
    if p_obj.text.strip() in HEADERS:
        p = p_obj._p
        prev = p.getprevious()
        # Skip if prev is not a paragraph (e.g. sectPr-only) or is the very first element
        if prev is None:
            continue
        if prev.tag != f'{Wq}p':
            continue
        if not is_blank_para(prev):
            # Insert blank before this header
            blank = make_blank_para()
            p.addprevious(blank)
            print(f"Inserted blank before {p_obj.text.strip()}")
        else:
            # Check that there's only ONE blank (remove extra blanks)
            prev2 = prev.getprevious()
            while prev2 is not None and prev2.tag == f'{Wq}p' and is_blank_para(prev2):
                to_del = prev2
                prev2 = prev2.getprevious()
                to_del.getparent().remove(to_del)
                print(f"Removed extra blank before {p_obj.text.strip()}")

doc.save('public/resume_2026.docx')
print("Saved docx")

from docx2pdf import convert
convert('public/resume_2026.docx', 'private-docs/resume_2026.pdf')
print("Saved PDF")
