"""Final consistency pass on resume_2026.docx:
1. Replace " — " (em dash) contextually (no more em dashes).
2. Strip rFonts and color overrides on all runs → use Normal style (single font, black).
3. Force black on Normal style.
4. Replace "GPT 5.4 (text) and Nano Banana (image)" with "text and image".
5. Remove "BLE" references.
"""
from docx import Document
from lxml import etree
import zipfile, shutil, re, os

DOCX = 'public/resume_2026.docx'
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Map of exact text → replacement
REPLACEMENTS = [
    # GPT/Banana
    ('GPT 5.4 (text) and Nano Banana (image)', 'text and image'),
    # P-values (remove from CV)
    (' p=.008 / p=.032 on verbal interaction.', ' on verbal interaction.'),
    ('p=.008 / p=.032 on verbal interaction.', 'on verbal interaction.'),
    ('p=.008 / p=.032', ''),
    # BLE
    ('personalized LLM agents over BLE', 'personalized LLM agents'),
    ('agent-to-agent reasoning over BLE', 'agent-to-agent reasoning'),
    (' over BLE', ''),
    (' BLE ', ' '),
    # Em dashes in specific phrases (most specific first)
    ('Visiting Research Student — MIT Media Lab', 'Visiting Research Student, MIT Media Lab'),
    ('Interaction Intelligence — Teleabsence', 'Interaction Intelligence: Teleabsence'),
    ('Winner, Connect track — MIT Hard Mode', 'Winner, Connect track at MIT Hard Mode'),
    ('Ph.D. Researcher — Institute for Future Technologies', 'Ph.D. Researcher, Institute for Future Technologies'),
    ('Tangible Co-Ideation — multi-agent LLM system', 'Tangible Co-Ideation: multi-agent LLM system'),
    ('IPheromone — personalized LLM agents', 'IPheromone: personalized LLM agents'),
    ('full PyTorch pipeline — dataset', 'full PyTorch pipeline (dataset'),
    ('training loop, evaluation', 'training loop, evaluation)'),
    ('mixed-methods evaluation — 3-arm', 'mixed-methods evaluation: 3-arm'),
    ('English (fluent — academic writing', 'English (fluent, academic writing'),
    ('Chess instructor) — French Chess Federation', 'Chess instructor, French Chess Federation'),
    ('DIFF certificate (chess instructor) — French Chess Federation', 'DIFF certificate (French Chess Federation)'),
    ('2nd place — PACA Open', '2nd place, PACA Open'),
    ('1st place — National IV Team Competition', '1st place, National IV Team Competition'),
    ('Eagles — ESILV US Football Team', 'Eagles, ESILV US Football Team'),
    ('University French Champion (2017) — US Football Association', 'University French Champion (2017), US Football Association'),
    # Catch-all em dash with surrounding spaces → comma+space (last resort)
    (' — ', ', '),
    ('—', '-'),
]

doc = Document(DOCX)

# Apply text replacements: walk all <w:t> elements and replace across concatenated text
# Strategy: for each paragraph, concatenate text from all runs, apply replacements,
# then redistribute proportionally — too risky. Better: replace within each <w:t>,
# then handle cross-run em-dashes by joining the runs first if needed.

def replace_in_paragraph_text(p):
    # Get all <w:t> nodes
    ts = p._p.findall(f'.//{W}t')
    full = ''.join(t.text or '' for t in ts)
    new = full
    for old, new_str in REPLACEMENTS:
        new = new.replace(old, new_str)
    if new == full:
        return False
    # Redistribute: put all new text in the first <w:t>, clear others
    if ts:
        ts[0].text = new
        ts[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        for t in ts[1:]:
            t.text = ''
    return True

changed_count = 0
for p in doc.paragraphs:
    if replace_in_paragraph_text(p):
        changed_count += 1
print(f"Replaced text in {changed_count} paragraphs")

# Strip rFonts and color overrides from all runs
stripped_fonts = 0
stripped_colors = 0
for p in doc.paragraphs:
    for r in p._p.findall(f'.//{W}r'):
        rPr = r.find(f'{W}rPr')
        if rPr is None:
            continue
        rf = rPr.find(f'{W}rFonts')
        if rf is not None:
            rPr.remove(rf)
            stripped_fonts += 1
        c = rPr.find(f'{W}color')
        if c is not None:
            rPr.remove(c)
            stripped_colors += 1
print(f"Stripped {stripped_fonts} rFonts, {stripped_colors} color overrides")

doc.save(DOCX)

# Now patch styles.xml: force black color on Normal style, single font (Trebuchet MS)
# Replace the docDefaults rFonts and color
with zipfile.ZipFile(DOCX, 'r') as z:
    files = {n: z.read(n) for n in z.namelist()}

styles = files['word/styles.xml'].decode('utf-8')

# Force docDefaults color to black 000000
styles = re.sub(
    r'<w:color w:val="[0-9A-Fa-f]{6}"[^/]*/>',
    '<w:color w:val="000000"/>',
    styles,
    count=1  # only first one (docDefaults)
)
# Also nuke any leftover color elements that aren't black
styles = re.sub(
    r'<w:color w:val="(?!000000")[0-9A-Fa-f]{6}"[^/]*/>',
    '<w:color w:val="000000"/>',
    styles
)
# Force docDefaults rFonts to Trebuchet MS
styles = re.sub(
    r'<w:rFonts[^/]*/>',
    '<w:rFonts w:ascii="Trebuchet MS" w:hAnsi="Trebuchet MS" w:cs="Trebuchet MS" w:eastAsia="Trebuchet MS"/>',
    styles
)

files['word/styles.xml'] = styles.encode('utf-8')

with zipfile.ZipFile(DOCX, 'w', zipfile.ZIP_DEFLATED) as z:
    for n, data in files.items():
        z.writestr(n, data)

print("Patched styles.xml: black color + Trebuchet MS font everywhere")

# Regenerate PDF
from docx2pdf import convert
convert(DOCX, 'private-docs/resume_2026.pdf')
print("Saved PDF")
