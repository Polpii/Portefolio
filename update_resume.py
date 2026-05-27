"""
Rebuild the PUBLICATIONS section of resume_2026.docx with a structured,
status-categorized layout (First-Author + Co-Authored, subgrouped by status).

- Removes the existing 3 publication bullets after the PUBLICATIONS heading
- Inserts new entries grouped by:
    * First-Author (Peer-Reviewed / Under Review / In Preparation)
    * Co-Authored  (Peer-Reviewed / Under Review / In Preparation)
- Bold = author's own name (Arslan, P.-P. with all italics for asterisks)
- Italics = venue name
- 9pt body, Calibri, * bullet (\u2022)
- Then regenerates the PDF via docx2pdf.
"""

from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOCX = "public/resume_2026.docx"
PDF_OUT = "private-docs/resume_2026.pdf"

BULLET = "\u2022"
FONT = "Calibri"
SIZE = Pt(9)
SUBHEADER_SIZE = Pt(9)

# (kind, text-segments) — segments are list of (text, {bold, italic})
# Entry format: list of segments forming one bullet line (without bullet prefix).
PUBS = {
    "First Author": {
        "Peer-Reviewed": [
            [
                ("Arslan, P.-P.", {"bold": True}),
                (", Xiao, X., et al. ReTouche: Embodied Representations for Self-Directed Piano Learning. ", {}),
                ("ACM CHI 2026", {"italic": True}),
                (" (full paper, accepted). [DOI: 10.1145/3772318.3791044]", {}),
            ],
        ],
        "Under Review": [
            [
                ("Li, Y. L.*, Kuang, Q.*, ", {}),
                ("Arslan, P.-P.*", {"bold": True}),
                (", Xiao, X., Ishii, H. Tangible Prompting: A Physical Vocabulary for Navigating Conceptual Spaces in Human-AI Co-Ideation. Submitted to ", {}),
                ("ACM UIST 2026", {"italic": True}),
                ("; Demo accepted at ", {}),
                ("ACM CHI 2026", {"italic": True}),
                (". (*equal contribution)", {}),
            ],
            [
                ("Li, Y. L.*, Kuang, Q.*, ", {}),
                ("Arslan, P.-P.*", {"bold": True}),
                (", Xiao, X., Ishii, H. Tangible Co-Ideation: Designing Embodied Prompting for Creative Thinking with Large Language Models. ", {}),
                ("ACM DIS 2026 SDC", {"italic": True}),
                (" (finalist). (*equal contribution) [PDF: paulpeterarslan.com/api/doc/Tangible/tangible_co_ideation_submission.pdf]", {}),
            ],
            [
                ("Arslan, P.-P.", {"bold": True}),
                (", et al. A Novel Method for Rhythmic Imitation of Finger Movements and the Effects of Auditory Stimuli and Feedback. Submitted to ", {}),
                ("Scientific Reports", {"italic": True}),
                (" (Springer Nature). [PDF: paulpeterarslan.com/api/doc/RhythmKaraoke/Rhythm%20Karaoke_22july2025.pdf]", {}),
            ],
        ],
        "In Preparation": [
            [
                ("Arslan, P.-P.", {"bold": True}),
                (", et al. Predicting Continuous Finger Forces from High-Density EMG with AI Regression Models. In preparation (target: ", {}),
                ("Nature", {"italic": True}),
                ("-family journal).", {}),
            ],
            [
                ("Arslan, P.-P.", {"bold": True}),
                (", et al. Rhythmic Finger Tapping Performance in Stroke Survivors: Effects of Auditory Stimuli and Feedback. In preparation.", {}),
            ],
        ],
    },
    "Co-Authored": {
        "Peer-Reviewed": [
            [
                ("Dziezuk, E., ", {}),
                ("Arslan, P.-P.", {"bold": True}),
                (", et al. (2025). M\u00e9canismes sensoriels de la r\u00e9cup\u00e9ration de la dext\u00e9rit\u00e9 manuelle apr\u00e8s AVC : protocole d\u2019une \u00e9tude de cohorte prospective. ", {}),
                ("Kin\u00e9sith\u00e9rapie Scientifique", {"italic": True}),
                (".", {}),
            ],
        ],
        "Under Review": [
            [
                ("Dziezuk, E., Badr, L., ", {}),
                ("Arslan, P.-P.", {"bold": True}),
                (", et al. A novel haptic stimulation device to improve finger movements \u2013 a validation and reliability study in healthy subjects. Submitted to ", {}),
                ("IEEE Transactions on Haptics", {"italic": True}),
                (" (Feb. 2026).", {}),
            ],
        ],
        "In Preparation": [
            [
                ("Duan, Z., ", {}),
                ("Arslan, P.-P.", {"bold": True}),
                (", Plantin, J., Lindberg, P. G., Wang, R. Motor Unit Coherence and Discharge Behavior During Single-Finger Isometric Contractions After Stroke. In preparation.", {}),
            ],
        ],
    },
}


def find_para(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i, p
    raise RuntimeError("paragraph not found")


def make_para(anchor, runs_data, *, bold_all=False, italic_all=False, size=SIZE, bullet=True, indent_left=None):
    """Insert a new paragraph immediately AFTER `anchor` paragraph, matching its style."""
    new_p = deepcopy(anchor._element)
    # Strip all runs in clone
    for r in list(new_p.findall(qn("w:r"))):
        new_p.remove(r)
    # Strip hyperlinks if any
    for h in list(new_p.findall(qn("w:hyperlink"))):
        new_p.remove(h)
    anchor._element.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)

    prefix = (BULLET + "  ") if bullet else ""
    if prefix:
        run = para.add_run(prefix)
        run.font.name = FONT
        run.font.size = size

    for text, props in runs_data:
        run = para.add_run(text)
        run.font.name = FONT
        run.font.size = size
        run.bold = bold_all or props.get("bold", False)
        run.italic = italic_all or props.get("italic", False)
    return para


def make_subheader(anchor, text):
    new_p = deepcopy(anchor._element)
    for r in list(new_p.findall(qn("w:r"))):
        new_p.remove(r)
    for h in list(new_p.findall(qn("w:hyperlink"))):
        new_p.remove(h)
    anchor._element.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = SUBHEADER_SIZE
    run.bold = True
    run.italic = True
    return para


def remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


def main():
    doc = Document(DOCX)

    # 1) Locate PUBLICATIONS header and EDUCATION header (boundary of the section)
    pub_idx, pub_header = find_para(doc, lambda t: t.strip() == "PUBLICATIONS")
    edu_boundary_idx, _ = find_para(doc, lambda t: t.strip() == "EDUCATION")

    # 2) Delete every paragraph strictly between pub header and EDUCATION header
    paragraphs = doc.paragraphs
    to_remove = paragraphs[pub_idx + 1 : edu_boundary_idx]
    for p in to_remove:
        remove_paragraph(p)

    # 3) Insert new entries AFTER pub_header, in reverse so that order is preserved.
    # We build a flat ordered list and insert each one after the previous.
    blocks = []  # (kind: 'sub'|'bullet', payload)
    for section, statuses in PUBS.items():
        blocks.append(("sub", section))
        for status, entries in statuses.items():
            blocks.append(("sub_inline", status))
            for entry in entries:
                blocks.append(("bullet", entry))

    anchor = pub_header
    for kind, payload in blocks:
        if kind == "sub":
            anchor = make_subheader(anchor, payload)
        elif kind == "sub_inline":
            # smaller italic status label, no bold
            new_p = deepcopy(anchor._element)
            for r in list(new_p.findall(qn("w:r"))):
                new_p.remove(r)
            for h in list(new_p.findall(qn("w:hyperlink"))):
                new_p.remove(h)
            anchor._element.addnext(new_p)
            para = Paragraph(new_p, anchor._parent)
            run = para.add_run(payload)
            run.font.name = FONT
            run.font.size = Pt(8.5)
            run.italic = True
            run.bold = False
            anchor = para
        else:
            anchor = make_para(anchor, payload)

    # 4) Remove any PhD/UPC entries accidentally left in EDUCATION (cleanup)
    edu_idx, _edu_header = find_para(doc, lambda t: t.strip() == "EDUCATION")
    while True:
        candidate = doc.paragraphs[edu_idx + 1]
        if candidate.text.startswith("PhD") or candidate.text.startswith("Universit"):
            remove_paragraph(candidate)
        else:
            break

    # 5) Update Ph.D. Researcher title line + supervisors date (bold + expected)
    _, researcher_para = find_para(doc, lambda t: t.startswith("Ph.D. Researcher"))
    researcher_para.runs[0].text = (
        "Ph.D. Researcher \u2014 Data Science & LLM Systems"
        " \u00b7 IFT / Neuroscience Institute Paris"
    )

    _, supervisors_para = find_para(doc, lambda t: "Supervisors:" in t and "Lindberg" in t)
    date_run = supervisors_para.runs[1]
    date_run.text = date_run.text.replace("Present", "Expected: July 2, 2026")
    date_run.bold = True

    doc.save(DOCX)
    print(f"Saved {DOCX}")

    # 6) Regenerate PDF
    from docx2pdf import convert
    convert(DOCX, PDF_OUT)
    print(f"Saved {PDF_OUT}")


if __name__ == "__main__":
    main()
